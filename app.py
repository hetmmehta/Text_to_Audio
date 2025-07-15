import os
import tempfile
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
from gtts import gTTS
import io
import base64
from ebooklib import epub
import traceback
import logging

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['AUDIO_FOLDER'] = 'audio'

# Ensure directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['AUDIO_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc', 'epub'}

# gTTS voice options (limited by language/accent)
VOICE_OPTIONS = [
    {"label": "English (US, Female)", "value": "en"},
    {"label": "English (UK, Female)", "value": "en-uk"},
    {"label": "English (US, Male)", "value": "en-us"},
    {"label": "English (UK, Male)", "value": "en-uk-rp"},
    {"label": "English (AU, Female)", "value": "en-au"},
    {"label": "English (IN, Female)", "value": "en-in"},
    {"label": "English (UK, North)", "value": "en-uk-north"},
    {"label": "English (UK, Scotland)", "value": "en-scotland"},
    {"label": "English (UK, London)", "value": "en-uk-london"},
    {"label": "English (UK, Leeds)", "value": "en-uk-leeds"},
    {"label": "English (UK, Manchester)", "value": "en-uk-manc"},
    {"label": "English (UK, Lancashire)", "value": "en-uk-lancashire"},
    {"label": "English (UK, Liverpool)", "value": "en-uk-liverpool"},
    {"label": "English (UK, Geordie)", "value": "en-uk-geordie"},
]

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file_path, file_extension):
    """Extract text from different file formats"""
    text = ""
    try:
        logger.info(f"Extracting text from {file_path} with extension {file_extension}")
        
        if file_extension == 'txt':
            # Try different encodings for text files
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    logger.info(f"Successfully read text file with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    continue
            if not text:
                return "Error extracting text: Could not decode text file with any supported encoding"
                
        elif file_extension == 'pdf':
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                logger.info(f"PDF has {len(pdf_reader.pages)} pages")
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    logger.info(f"Extracted {len(page_text)} characters from page {i+1}")
                    
        elif file_extension in ['docx', 'doc']:
            doc = Document(file_path)
            logger.info(f"Document has {len(doc.paragraphs)} paragraphs")
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    
        elif file_extension == 'epub':
            book = epub.read_epub(file_path)
            logger.info(f"EPUB has {len(book.get_items())} items")
            for item in book.get_items():
                if item.get_type() == epub.ITEM_DOCUMENT:
                    try:
                        from bs4 import BeautifulSoup
                        soup = BeautifulSoup(item.get_content(), 'html.parser')
                        item_text = ' '.join([t for t in soup.stripped_strings])
                        if item_text:
                            text += item_text + "\n"
                    except Exception as e:
                        logger.error(f"Error parsing EPUB item: {e}")
                        continue
                        
        text = text.strip()
        logger.info(f"Extracted {len(text)} characters of text")
        return text if text else "Error extracting text: No text content found"
        
    except Exception as e:
        error_msg = f"Error extracting text: {str(e)}"
        logger.error(f"{error_msg}\n{traceback.format_exc()}")
        return error_msg

def text_to_speech(text, language='en'):
    """Convert text to speech and return audio data"""
    try:
        logger.info(f"Converting text to speech with language: {language}")
        tts = gTTS(text=text, lang=language, tld='com', slow=False)
        audio_buffer = io.BytesIO()
        tts.write_to_fp(audio_buffer)
        audio_buffer.seek(0)
        logger.info("Successfully converted text to speech")
        return audio_buffer
    except Exception as e:
        logger.error(f"Error in text_to_speech: {e}\n{traceback.format_exc()}")
        return None

def get_voice_params(voice_value):
    # Map UI voice value to gTTS language/accent
    mapping = {
        'en': ('en', 'com'),
        'en-uk': ('en', 'co.uk'),
        'en-us': ('en', 'com'),
        'en-au': ('en', 'com.au'),
        'en-in': ('en', 'co.in'),
        'en-uk-rp': ('en', 'co.uk'),
        'en-uk-north': ('en', 'co.uk'),
        'en-scotland': ('en', 'co.uk'),
        'en-uk-london': ('en', 'co.uk'),
        'en-uk-leeds': ('en', 'co.uk'),
        'en-uk-manc': ('en', 'co.uk'),
        'en-uk-lancashire': ('en', 'co.uk'),
        'en-uk-liverpool': ('en', 'co.uk'),
        'en-uk-geordie': ('en', 'co.uk'),
    }
    return mapping.get(voice_value, ('en', 'com'))

def text_to_speech_with_voice(text, voice_value):
    lang, tld = get_voice_params(voice_value)
    try:
        logger.info(f"Converting text to speech with voice: {voice_value} (lang={lang}, tld={tld})")
        tts = gTTS(text=text, lang=lang, tld=tld, slow=False)
        audio_buffer = io.BytesIO()
        tts.write_to_fp(audio_buffer)
        audio_buffer.seek(0)
        logger.info("Successfully converted text to speech with voice")
        return audio_buffer
    except Exception as e:
        logger.error(f"Error in text_to_speech_with_voice: {e}\n{traceback.format_exc()}")
        return None

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/voice-options')
def voice_options():
    return jsonify(VOICE_OPTIONS)

@app.route('/convert', methods=['POST'])
def convert_to_audio():
    try:
        logger.info("Received file conversion request")
        
        if 'file' not in request.files:
            logger.error("No file uploaded")
            return jsonify({'error': 'No file uploaded'}), 400
            
        file = request.files['file']
        if file.filename == '':
            logger.error("No file selected")
            return jsonify({'error': 'No file selected'}), 400
            
        if not allowed_file(file.filename):
            logger.error(f"File type not supported: {file.filename}")
            return jsonify({'error': 'File type not supported'}), 400
            
        voice = request.form.get('voice', 'en')
        filename = secure_filename(file.filename)
        file_extension = filename.rsplit('.', 1)[1].lower()
        
        logger.info(f"Processing file: {filename} (extension: {file_extension}, voice: {voice})")
        
        # Save file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_extension}') as temp_file:
            file.save(temp_file.name)
            temp_file_path = temp_file.name
            
        logger.info(f"File saved temporarily at: {temp_file_path}")
        
        # Extract text
        extracted_text = extract_text_from_file(temp_file_path, file_extension)
        
        # Clean up temp file
        try:
            os.unlink(temp_file_path)
        except:
            pass
            
        if not extracted_text or extracted_text.startswith("Error"):
            logger.error(f"Text extraction failed: {extracted_text}")
            return jsonify({'error': extracted_text or 'No text could be extracted'}), 400
            
        # Convert to speech
        audio_buffer = text_to_speech_with_voice(extracted_text, voice)
        if audio_buffer is None:
            logger.error("Failed to convert text to speech")
            return jsonify({'error': 'Failed to convert text to speech'}), 500
            
        audio_data = base64.b64encode(audio_buffer.getvalue()).decode('utf-8')
        
        logger.info("Successfully converted file to audio")
        return jsonify({
            'success': True,
            'text': extracted_text,
            'audio_data': audio_data,
            'filename': filename
        })
        
    except Exception as e:
        error_msg = f'An error occurred: {str(e)}'
        logger.error(f"{error_msg}\n{traceback.format_exc()}")
        return jsonify({'error': error_msg}), 500

@app.route('/convert-text', methods=['POST'])
def convert_text_to_audio():
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        voice = data.get('voice', 'en')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
            
        audio_buffer = text_to_speech_with_voice(text, voice)
        if audio_buffer is None:
            return jsonify({'error': 'Failed to convert text to speech'}), 500
            
        audio_data = base64.b64encode(audio_buffer.getvalue()).decode('utf-8')
        return jsonify({
            'success': True,
            'audio_data': audio_data
        })
    except Exception as e:
        error_msg = f'An error occurred: {str(e)}'
        logger.error(f"{error_msg}\n{traceback.format_exc()}")
        return jsonify({'error': error_msg}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)